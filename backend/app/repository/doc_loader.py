import os
from docx import Document
from llama_index.core.schema import Document as LlamaDocument

def load_docx_file(filepath: str) -> list[LlamaDocument]:
    doc = Document(filepath)
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    metadata = {"file_name": os.path.basename(filepath)}
    return [LlamaDocument(text=full_text, metadata=metadata)]